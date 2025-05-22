import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document

class AdvancedTableEditorApp:
    def __init__(self, master):
        self.master = master
        master.title("Word 표 편집기 (고급형)")

        self.doc = None
        self.file_path = None

        tk.Button(master, text="Word 파일 열기", command=self.load_docx).pack(pady=5)

        self.info_text = tk.Text(master, height=10, width=70)
        self.info_text.pack(pady=5)

        # 셀 미리보기
        preview_frame = tk.Frame(master)
        preview_frame.pack()
        self.t_idx = tk.StringVar()
        self.r_idx = tk.StringVar()
        self.c_idx = tk.StringVar()
        self.preview_result = tk.StringVar()

        self._add_labeled_entry(preview_frame, "테이블 번호:", self.t_idx)
        self._add_labeled_entry(preview_frame, "행 번호:", self.r_idx)
        self._add_labeled_entry(preview_frame, "열 번호:", self.c_idx)

        tk.Button(master, text="셀 내용 미리보기", command=self.preview_cell).pack(pady=3)
        tk.Label(master, textvariable=self.preview_result, fg="blue").pack()

        # 일괄 입력
        tk.Label(master, text="여러 셀 입력 (예: 0,1,2=홍길동)").pack()
        self.bulk_text = tk.Text(master, height=6, width=70)
        self.bulk_text.pack(pady=3)

        tk.Button(master, text="일괄 적용", command=self.apply_bulk).pack(pady=5)
        tk.Button(master, text="저장하기", command=self.save_docx).pack(pady=5)

    def _add_labeled_entry(self, parent, label, variable):
        frame = tk.Frame(parent)
        frame.pack(side=tk.LEFT, padx=2)
        tk.Label(frame, text=label).pack()
        tk.Entry(frame, textvariable=variable, width=5).pack()

    def load_docx(self):
        self.file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if self.file_path:
            try:
                self.doc = Document(self.file_path)
                info = [f"총 테이블 수: {len(self.doc.tables)}"]
                for i, table in enumerate(self.doc.tables):
                    rows = len(table.rows)
                    cols = len(table.rows[0].cells) if rows > 0 else 0
                    info.append(f" - 테이블 {i}: {rows}행 × {cols}열")
                self.info_text.delete("1.0", tk.END)
                self.info_text.insert(tk.END, "\n".join(info))
                messagebox.showinfo("성공", "Word 문서를 불러왔습니다.")
            except Exception as e:
                messagebox.showerror("오류", f"문서 로드 실패: {e}")

    def preview_cell(self):
        if not self.doc:
            messagebox.showerror("오류", "문서를 먼저 여세요.")
            return
        try:
            t = int(self.t_idx.get())
            r = int(self.r_idx.get())
            c = int(self.c_idx.get())
            cell_text = self.doc.tables[t].cell(r, c).text
            self.preview_result.set(f"셀 내용: {cell_text}")
        except Exception as e:
            self.preview_result.set(f"에러: {e}")

    def apply_bulk(self):
        if not self.doc:
            messagebox.showerror("오류", "문서를 먼저 여세요.")
            return
        try:
            lines = self.bulk_text.get("1.0", tk.END).strip().splitlines()
            count = 0
            for line in lines:
                if '=' not in line:
                    continue
                left, value = line.split('=', 1)
                t, r, c = map(int, left.strip().split(','))
                self.doc.tables[t].cell(r, c).text = value.strip()
                count += 1
            messagebox.showinfo("적용 완료", f"{count}개의 셀을 수정했습니다.")
        except Exception as e:
            messagebox.showerror("오류", f"적용 실패: {e}")

    def save_docx(self):
        if not self.doc:
            messagebox.showerror("오류", "문서를 먼저 여세요.")
            return
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if save_path:
            self.doc.save(save_path)
            messagebox.showinfo("저장 완료", "문서가 저장되었습니다.")

if __name__ == "__main__":
    root = tk.Tk()
    app = AdvancedTableEditorApp(root)
    root.mainloop()
